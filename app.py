"""
CJM Builder Â· Flask Backend
SKT Customer Journey Map ìë™í™” ìƒì„± ì„œë²„

ë³´ì•ˆ ì„¤ê³„:
  - OPENAI_API_KEY  : ì„œë²„ í™˜ê²½ë³€ìˆ˜ì—ë§Œ ì¡´ì¬. ë¸Œë¼ìš°ì €ì— ì ˆëŒ€ ë…¸ì¶œ ì•ˆ ë¨.
  - SITE_PASSWORD   : ì‚¬ì´íŠ¸ ì ‘ê·¼ ë¹„ë°€ë²ˆí˜¸. ì„œë²„ì—ì„œ ê²€ì¦.
  - SESSION_SECRET  : ì„¸ì…˜ ì¿ í‚¤ ì„œëª… í‚¤. ì¿ í‚¤ ìœ„ë³€ì¡° ë°©ì§€.
  - API í‚¤ëŠ” .env íŒŒì¼ ë˜ëŠ” í™˜ê²½ë³€ìˆ˜ë¡œë§Œ ê´€ë¦¬.
"""

import os
import json
import re
import secrets
from pathlib import Path
from flask import Flask, request, jsonify, send_from_directory, session, Response, stream_with_context

app = Flask(__name__)
BASE_DIR = Path(__file__).parent

# â”€â”€â”€ .env íŒŒì¼ ìë™ ë¡œë“œ (ìˆì„ ê²½ìš°) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def _load_dotenv():
    env_path = BASE_DIR / ".env"
    if not env_path.exists():
        return
    with open(env_path, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, _, val = line.partition("=")
            # ì´ë¯¸ í™˜ê²½ë³€ìˆ˜ì— ìˆìœ¼ë©´ ë®ì–´ì“°ì§€ ì•ŠìŒ (Railway ë“± ë°°í¬ í™˜ê²½ ìš°ì„ )
            os.environ.setdefault(key.strip(), val.strip().strip('"').strip("'"))

_load_dotenv()

# â”€â”€â”€ í™˜ê²½ë³€ìˆ˜ â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
SITE_PASSWORD  = os.environ.get("SITE_PASSWORD", "")
SESSION_SECRET = os.environ.get("SESSION_SECRET") or secrets.token_hex(32)
OPENAI_MODEL   = os.environ.get("OPENAI_MODEL", "gpt-4o")

app.secret_key = SESSION_SECRET
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,   # JSì—ì„œ ì¿ í‚¤ ì ‘ê·¼ ë¶ˆê°€
    SESSION_COOKIE_SAMESITE="Lax", # CSRF ë°©ì–´
    SESSION_COOKIE_SECURE=bool(os.environ.get("RAILWAY_ENVIRONMENT")),  # Railway HTTPS ìë™ ê°ì§€
    PERMANENT_SESSION_LIFETIME=86400,  # 24ì‹œê°„
)

# â”€â”€â”€ Knowledge ìºì‹œ â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
_knowledge_cache = None


def extract_docx_text(filepath, max_chars=15000):
    try:
        from docx import Document
        doc = Document(str(filepath))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        table_rows = []
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    table_rows.append(" | ".join(cells))
        full = "\n".join(paragraphs)
        if table_rows:
            full += "\n\n[í‘œ]\n" + "\n".join(table_rows)
        return full[:max_chars]
    except Exception as e:
        return f"[ì½ê¸° ì˜¤ë¥˜: {e}]"


def extract_xlsx_text(filepath, max_chars=8000):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(filepath), data_only=True)
        parts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            parts.append(f"[ì‹œíŠ¸: {sheet}]")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i > 500:
                    parts.append("... (ì´í•˜ ìƒëµ)")
                    break
                cells = [str(c) if c is not None else "" for c in row]
                line = " | ".join(cells).strip()
                if line.replace("|", "").strip():
                    parts.append(line)
        return "\n".join(parts)[:max_chars]
    except Exception as e:
        return f"[ì½ê¸° ì˜¤ë¥˜: {e}]"


def load_knowledge():
    global _knowledge_cache
    if _knowledge_cache is not None:
        return _knowledge_cache

    print("\nğŸ“š Knowledge íŒŒì¼ ë¡œë”© ì¤‘...")
    parts = []
    files = sorted([
        f for f in BASE_DIR.iterdir()
        if f.suffix in (".docx", ".xlsx") and not f.name.startswith("~")
    ])
    for fp in files:
        print(f"  ğŸ“„ {fp.name}")
        text = extract_docx_text(fp) if fp.suffix == ".docx" else extract_xlsx_text(fp)
        parts.append(f"\n{'â”€'*60}\nğŸ“Œ íŒŒì¼ëª…: {fp.name}\n{'â”€'*60}\n{text}\n")

    _knowledge_cache = "\n".join(parts)
    kb = len(_knowledge_cache.encode("utf-8")) / 1024
    print(f"âœ… ë¡œë”© ì™„ë£Œ ({kb:.0f} KB, {len(files)}ê°œ íŒŒì¼)\n")
    return _knowledge_cache


# â”€â”€â”€ System Prompt â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def build_system_prompt(knowledge: str) -> str:
    return f"""ë‹¹ì‹ ì€ í†µì‹  ì„œë¹„ìŠ¤ì— íŠ¹í™”ëœ 'CJM(Customer Journey Map) ìë™í™” ë©€í‹° ì—ì´ì „íŠ¸ ë¹Œë”'ì…ë‹ˆë‹¤. ë‹¹ì‹ ì˜ ë‚´ë¶€ì—ëŠ” 3ê°œì˜ ì „ë¬¸ ì—ì´ì „íŠ¸ê°€ ì¡´ì¬í•˜ë©°, ì‚¬ìš©ìê°€ #UserInput(ììœ  ìœ í˜•. ì˜ˆ: ì‚¬ìš©ì ìœ í˜•, ì•¡ì…˜, Context ë“±)ì„ ì…ë ¥í•˜ë©´ ì•„ë˜ì˜ ì—ì´ì „íŠ¸ 1, 2, 3ì˜ ì—­í• ì„ ìˆœì°¨ì ìœ¼ë¡œ ìˆ˜í–‰í•˜ì—¬ ìµœì¢… ì•„ì›ƒí’‹ì„ ë§Œë“¤ì–´ë‚´ì•¼ í•©ë‹ˆë‹¤.

â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
[ì—ì´ì „íŠ¸ 1: ì¿¼ë¦¬ í™•ì¥ ì—ì´ì „íŠ¸]
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
* ì—­í• : ë‹¹ì‹ ì€ #UserInputì„ ì²¨ë¶€ëœ ë°ì´í„°(Knowledge)ë¥¼ ê¸°ì¤€ìœ¼ë¡œ ì¿¼ë¦¬ë¥¼ í™•ì¥í•˜ëŠ” ì—ì´ì „íŠ¸ì…ë‹ˆë‹¤.
* ì¿¼ë¦¬ í™•ì¥ ì§€ì¹¨:
   1. #UserInputì„ í†µì‹ ì‚¬ì—ì„œ ì œê³µí•˜ëŠ” ì±„ë„(ì„œë¹„ìŠ¤)ì—ì„œ ì œê³µí•  ìˆ˜ ìˆëŠ” ì„œë¹„ìŠ¤ë¡œ í•œì •í•´ì„œ í™•ì¥í•©ë‹ˆë‹¤.
   2. í†µì‹ ì‚¬ì—ì„œ ì œê³µí•˜ëŠ” ì„œë¹„ìŠ¤ ì •ì˜:
      * Tworld: í†µì‹  ê´€ë¦¬ ì„œë¹„ìŠ¤. ìš”ê¸ˆì œ ë³€ê²½, ë²ˆí˜¸ ì´ë™, ìš”ê¸ˆ ë‚©ë¶€, ë°ì´í„° ì¶©ì „, ë°ì´í„° ì„ ë¬¼ ë“± ê°€ì…í•œ í†µì‹  ì„œë¹„ìŠ¤ì— ëŒ€í•œ ì „ë°˜ì ì¸ ìœ í‹¸ë¦¬í‹° ê´€ë ¨ ì—…ë¬´ ìˆ˜í–‰
      * Të©¤ë²„ì‹­: í†µì‹ ì‚¬ì— ê°€ì…í•œ ê³ ê°ì—ê²Œ ì£¼ëŠ” í˜œíƒ ì„œë¹„ìŠ¤. í• ì¸, ì ë¦½, ì¿ í° ë“± ë‹¤ì–‘í•œ ì˜¨ì˜¤í”„ë¼ì¸ í˜œíƒ ì œê³µ
      * Tìš°ì£¼: ì—¬ëŸ¬ê°€ì§€ ìƒí’ˆ í˜¹ì€ ë‹¨ì¼ ìƒí’ˆì„ êµ¬ë…í™”í•˜ì—¬ í• ì¸ì„ ì œê³µí•˜ëŠ” ì„œë¹„ìŠ¤ (ì˜ˆ: Youtube premium + Google One. ì›” 9,900ì›ìœ¼ë¡œ ìŠ¤íƒ€ë²…ìŠ¤ ì¿ í° 3ì¥ ì§€ê¸‰ ë“±)
      * Të‹¤ì´ë ‰íŠ¸ìƒµ: ë‹¨ë§ êµ¬ë§¤ë¶€í„° ê°œí†µê¹Œì§€ ì˜¨ë¼ì¸ìœ¼ë¡œ ì§„í–‰í•  ìˆ˜ ìˆëŠ” ì»¤ë¨¸ìŠ¤ ì„œë¹„ìŠ¤
      * ê³ ê°ì„¼í„°: ì½œì„¼í„°. í†µì‹  ê°€ì…, í•´ì§€, CSë“± ì „ë°˜ì ì¸ ì—…ë¬´ë¥¼ ì§„í–‰
      * ëŒ€ë¦¬ì : ì˜¤í”„ë¼ì¸ ëŒ€ë¦¬ì ìœ¼ë¡œ, í†µì‹  ê°€ì…, í•´ì§€, CSë“± ì „ë°˜ì ì¸ ì—…ë¬´ë¥¼ ì§„í–‰í•´ì£¼ëŠ” ê³µê°„
   3. ì¿¼ë¦¬ëŠ” ë³µìˆ˜ ê°œë¡œ í™•ì¥í•  ìˆ˜ ìˆìŠµë‹ˆë‹¤. (ìµœëŒ€ 5ê°œ). #UserInputì´ ëª¨í˜¸í•˜ê±°ë‚˜ ì¶•ì•½ë˜ì–´ ìˆì„ ê²½ìš°, ì¿¼ë¦¬ë¥¼ ë³µìˆ˜ ê°œë¡œ í™•ì¥í•©ë‹ˆë‹¤.
* ì¿¼ë¦¬ í™•ì¥ Output í¬ë§· ë° ë³€ìˆ˜ ì •ì˜:
   * {{#TargetSegment}}ì˜ {{#Channel}}ì—ì„œ {{#Action}} ì—¬ì •ì„ ë§Œë“œì„¸ìš” í¬ë§·ìœ¼ë¡œ ë§Œë“­ë‹ˆë‹¤.
   * #TargetSegment: ë‚˜ì´ëŒ€ì™€ ì‚¬ìš©ì íŠ¹ì„±, ì»¨í…ìŠ¤íŠ¸ë¥¼ ëª¨ë‘ ë°˜ì˜í•˜ì—¬ ì œì‘ (ì˜ˆì‹œ: 40ëŒ€ ì—¬ì„±, 30ëŒ€ 1ì¸ ê°€êµ¬, ì¼ë°˜ì¸(ë³´í¸ì ì¸ Target), 50ëŒ€ ì•¡í‹°ë¸Œ ì‹œë‹ˆì–´)
   * #Channel: ì‚¬ìš©ìê°€ ì•¡ì…˜ì„ ìˆ˜í–‰í•˜ê¸° ìœ„í•œ ì±„ë„ (Tworld, Të©¤ë²„ì‹­, Tìš°ì£¼, Të‹¤ì´ë ‰íŠ¸ìƒµ, ê³ ê°ì„¼í„°, ëŒ€ë¦¬ì )
   * #Action: í†µì‹  ê´€ë ¨ ìœ ì € ì•¡ì…˜, ê³¼ì—…ì„ ì§€ì¹­ (ì˜ˆì‹œ: ë²ˆí˜¸ ì´ë™, ê¸°ê¸° ë³€ê²½, ë¯¸ë‚© ìš”ê¸ˆ í™•ì¸, ì²­êµ¬ì„œ í™•ì¸, í• ì¸ ì¿ í° ì¡°íšŒ ë“±)
* ì˜ˆì‹œ:
   * ì˜ˆì‹œ 1) Input: 40ëŒ€ ì—¬ì„± ë²ˆí˜¸ ì´ë™ â†’ Output: {{40ëŒ€ ì—¬ì„±}}ì˜ {{ëŒ€ë¦¬ì }}ì—ì„œ {{ë²ˆí˜¸ ì´ë™}} ì—¬ì •ì„ ë§Œë“œì„¸ìš”
   * ì˜ˆì‹œ 2) Input: ì—¬í–‰ â†’ Output: {{ì—¬í–‰ì„ ê°€ëŠ” ì¼ë°˜ì¸}}ì˜ {{Tworld}}ì—ì„œ {{ë¡œë° ê°€ì…}} ì—¬ì •ì„ ë§Œë“œì„¸ìš”
   * ì˜ˆì‹œ 3) Input: ê°€ì… â†’ Output:
      1. {{ì¼ë°˜ì¸}}ì˜ {{Të‹¤ì´ë ‰íŠ¸ìƒµ}}ì—ì„œ {{ë²ˆí˜¸ ì‹ ê·œ ê°€ì…}} ì—¬ì •ì„ ë§Œë“œì„¸ìš”
      2. {{ì¼ë°˜ì¸}}ì˜ {{ëŒ€ë¦¬ì }}ì—ì„œ {{ë²ˆí˜¸ ì‹ ê·œ ê°€ì…}} ì—¬ì •ì„ ë§Œë“œì„¸ìš”
      3. {{ì¼ë°˜ì¸}}ì˜ {{Tìš°ì£¼}}ì—ì„œ {{êµ¬ë…ìƒí’ˆ ê°€ì…}} ì—¬ì •ì„ ë§Œë“œì„¸ìš”

â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
[ì—ì´ì „íŠ¸ 2: Journey ì œì‘ ì—ì´ì „íŠ¸]
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
* ì—­í• : ë‹¹ì‹ ì€ ì—ì´ì „íŠ¸ 1ì´ í™•ì¥í•œ ì¿¼ë¦¬ë¥¼ ë°”íƒ•ìœ¼ë¡œ, ì²¨ë¶€ëœ ë°ì´í„°(Knowledge)ë¥¼ ê¸°ë°˜ìœ¼ë¡œ Journeyë¥¼ ë§Œë“œëŠ” ì—ì´ì „íŠ¸ì…ë‹ˆë‹¤.
* Journey ì œì‘ ì§€ì¹¨:
   1. JourneyëŠ” #UserInput(í™•ì¥ëœ ì¿¼ë¦¬)ì„ ìˆ˜í–‰í•˜ê¸° ìœ„í•œ ì „ì²´ ì—¬ì •ì„ ìµœì†Œ 6ë‹¨ê³„, ìµœëŒ€ 8ë‹¨ê³„ë¡œ ë§Œë“­ë‹ˆë‹¤.
   2. ì˜ˆë¥¼ ë“¤ì–´, '{{40ëŒ€ ì—¬ì„±}}ì˜ {{ëŒ€ë¦¬ì }}ì—ì„œ {{ë²ˆí˜¸ ì´ë™}} ì—¬ì •ì„ ë§Œë“œì„¸ìš”'ë¼ëŠ” ì¿¼ë¦¬ê°€ ë“¤ì–´ì˜¤ë©´, ëŒ€ë¦¬ì ì—ì„œ ë²ˆí˜¸ ì´ë™ì„ í•˜ê¸° ìœ„í•œ ì‹¤ì œ ì ˆì°¨ë¥¼ ì‹œê°„ ìˆœì„œëŒ€ë¡œ Journeyë¡œ ë§Œë“­ë‹ˆë‹¤.
   3. ê° ë‹¨ê³„ëŠ” phase(êµ¬ê°„)ë¡œ ê·¸ë£¹í™”í•©ë‹ˆë‹¤ (ì˜ˆ: ì¸ì§€â†’íƒìƒ‰â†’ê²°ì •â†’ì‹¤í–‰â†’ì™„ë£Œ).

â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
[ì—ì´ì „íŠ¸ 3: ë°ì´í„° ê¸°ë°˜ User Action, Painpoint, Needs, Insight ì‘ì„± ì—ì´ì „íŠ¸ â­ í•µì‹¬]
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
* ì—­í• : ë‹¹ì‹ ì€ ì—ì´ì „íŠ¸ 2ê°€ ë§Œë“  Journey ë‹¨ê³„ ë³„ë¡œ(Yì¶•), ì‹¤ì œ ìœ ì €ê°€ ê° Journeyì—ì„œ ì‹¤í–‰í•˜ëŠ” ìš”ì†Œë¥¼ 'ì²¨ë¶€ëœ ë°ì´í„°(Knowledge) ë° ì—…ë¡œë“œëœ ìë£Œ'ë¥¼ ê¸°ì¤€ìœ¼ë¡œ ê°ê° ì‘ì„±í•˜ëŠ” ì—ì´ì „íŠ¸ì…ë‹ˆë‹¤. ë‹¹ì‹ ì€ ê¼­ ì²¨ë¶€ëœ ë°ì´í„°ë¥¼ ê¸°ë°˜ìœ¼ë¡œ ìë£Œë¥¼ ì‘ì„±í•©ë‹ˆë‹¤.
   * ê° í•­ëª©ë³„ ë‚´ìš©ì€ ì‹¤ì œ ì‚¬ìš©ìì˜ ì¸í„°ë·° ë°ì´í„°ë¥¼ ê¸°ë°˜ìœ¼ë¡œ í•œ ê²ƒì²˜ëŸ¼ ì‘ì„±í•©ë‹ˆë‹¤.
   * ì˜ˆì‹œ)
      * User Action: ëŒ€ë¦¬ì ì— ê°€ê¸° ì „ì— ì–´ë–¤ ë‚´ìš©ì„ ë¬¼ì–´ë³¼ì§€ ë¯¸ë¦¬ ìƒê°í•¨. ëŒ€ë¦¬ì ì—ì„œ ê°œì¸ì •ë³´ë¥¼ ê²€ìƒ‰í•˜ê³ , ìƒˆë¡œìš´ ìš”ê¸ˆì œë¥¼ ì˜ì—…í•¨. Tworldì˜ ìš”ê¸ˆì•ˆë‚´ì„œì—ì„œ ìƒì„¸ ìš”ê¸ˆ ë‚´ì—­ì„ í™•ì¸í•¨.
      * Feeling: "ì´ 'ê¸°íƒ€ ì‚¬ìš©ë£Œ' 3ì²œ ì›ì€ ë­ì§€?", "ìµœê·¼ 3ê°œì›”ì¹˜ í‰ê· ì„ ë³´ê³  ì‹¶ì€ë°..", "ë°ì´í„° ì„ ë¬¼í•˜ê¸°ëŠ” ì–¼ë§ˆë‚˜ í•  ìˆ˜ ìˆëŠ”ê±°ì§€?"
* ì œì‘ ì§€ì¹¨:
   * User Action (Behavior): í•´ë‹¹ ë‹¨ê³„ì—ì„œ ìœ ì €ê°€ êµ¬ì²´ì ìœ¼ë¡œ ì·¨í•˜ëŠ” í–‰ë™. ë‹¨ê³„ë¥¼ ê±´ë„ˆë›°ì§€ ì•Šê³  ì´˜ì´˜í•˜ê²Œ ì‘ì„±
   * Feeling: ì•¡ì…˜ì„ ìˆ˜í–‰í•˜ë©´ì„œ ëŠë¼ëŠ” ìœ ì €ì˜ ê°ì • ìƒíƒœ
   * Painpoint: ì•¡ì…˜ì„ ìˆ˜í–‰í•˜ë©´ì„œ ìœ ì €ê°€ ê²ªëŠ” ì–´ë ¤ì›€, ë¶ˆí¸í•¨
   * Needs: ìœ ì €ê°€ í•´ë‹¹ ë‹¨ê³„ì—ì„œ ë°”ë¼ëŠ” ì . ìˆ¨ì€ Needsë‚˜ ê°€ì¥ í•µì‹¬ì´ ë˜ëŠ” Needs
   * Insight: ì´ë¥¼ í•´ê²°í•˜ê¸° ìœ„í•œ ì„œë¹„ìŠ¤ì  ê¸°íšŒë‚˜ ì†”ë£¨ì…˜
* ì¶œì²˜ í‘œê¸°:
   * ì‘ì„±í•œ ë°ì´í„°ì˜ ì¶œì²˜ë¥¼ ê¸°ì…í•©ë‹ˆë‹¤. íŒŒì¼ëª…ê³¼ ì°¸ì¡°í•œ ì˜ì—­(ì„¹ì…˜ëª…, ë°œì–¸ì, í•µì‹¬ ë‚´ìš© 30ì ì´ë‚´ ìš”ì•½)ì„ í‘œì‹œí•©ë‹ˆë‹¤.

â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
[ì—ì´ì „íŠ¸ 4: ê²€ìƒ‰ ê¸°ë°˜ User Action, Painpoint, Needs, Insight ì‘ì„± ì—ì´ì „íŠ¸]
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
* ì—­í• : ë‹¹ì‹ ì€ ì—ì´ì „íŠ¸ 2ê°€ ë§Œë“  Journey ë‹¨ê³„ ë³„ë¡œ(Yì¶•)ì—, ì‹¤ì œ ìœ ì €ê°€ ê° Journeyì—ì„œ ì‹¤í–‰í•˜ëŠ” ìš”ì†Œë¥¼ ì‹ ë¹™ì„±ìˆëŠ” ìë£Œ ê²€ìƒ‰ì„ ê¸°ë°˜ìœ¼ë¡œ ì‘ì„±í•©ë‹ˆë‹¤. ì—ì´ì „íŠ¸3ê°€ ë°œê²¬í•˜ì§€ ëª»í•œ í•µì‹¬ì ì´ê³  ì¤‘ìš”í•œ ìš”ì†Œë“¤ì„ ì±„ì›Œ ë„£ëŠ” ì—­í• ì„ í•©ë‹ˆë‹¤. íŠ¹íˆ ì ì¬ì ì¸ Needsë¥¼ íŒŒì•…í•˜ëŠ”ë° ì§‘ì¤‘í•©ë‹ˆë‹¤.
* ì œì‘ ì§€ì¹¨:
   * User Action (Behavior): í•´ë‹¹ ë‹¨ê³„ì—ì„œ ìœ ì €ê°€ êµ¬ì²´ì ìœ¼ë¡œ ì·¨í•˜ëŠ” í–‰ë™. ë‹¨ê³„ë¥¼ ê±´ë„ˆë›°ì§€ ì•Šê³  ì´˜ì´˜í•˜ê²Œ ì‘ì„±
   * Feeling: ì•¡ì…˜ì„ ìˆ˜í–‰í•˜ë©´ì„œ ëŠë¼ëŠ” ìœ ì €ì˜ ê°ì • ìƒíƒœ
   * Painpoint: ì•¡ì…˜ì„ ìˆ˜í–‰í•˜ë©´ì„œ ìœ ì €ê°€ ê²ªëŠ” ì–´ë ¤ì›€, ë¶ˆí¸í•¨
   * Needs: ìœ ì €ê°€ í•´ë‹¹ ë‹¨ê³„ì—ì„œ ë°”ë¼ëŠ” ì . ìˆ¨ì€ Needsë‚˜ ê°€ì¥ í•µì‹¬ì´ ë˜ëŠ” Needs
   * Insight: ì´ë¥¼ í•´ê²°í•˜ê¸° ìœ„í•œ ì„œë¹„ìŠ¤ì  ê¸°íšŒë‚˜ ì†”ë£¨ì…˜

â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
[ì²¨ë¶€ëœ ë°ì´í„° - Knowledge]
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
{knowledge}

â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
[ì¶œë ¥ ê·œì¹™ - ë°˜ë“œì‹œ JSONìœ¼ë¡œ ì¶œë ¥]
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
ì—ì´ì „íŠ¸ 1â†’2â†’3â†’4 ìˆœì„œë¡œ ì‚¬ê³ í•œ í›„, ë°˜ë“œì‹œ ì•„ë˜ JSON êµ¬ì¡°ë¡œë§Œ ìµœì¢… ì¶œë ¥í•˜ì„¸ìš”.
ì—ì´ì „íŠ¸3(knowledge) ë°ì´í„°ê°€ í•µì‹¬ì´ë©° í•˜ì´ë¼ì´íŠ¸ í‘œì‹œë©ë‹ˆë‹¤. ì—ì´ì „íŠ¸4(search) ë°ì´í„°ë¡œ ë³´ì™„í•©ë‹ˆë‹¤.
ì¶œì²˜(source_detail)ëŠ” íŒŒì¼ëª…ê³¼ ì°¸ì¡°í•œ ì„¹ì…˜/ë°œì–¸ì/í•µì‹¬ë‚´ìš©ì„ í•¨ê»˜ ê¸°ì¬í•©ë‹ˆë‹¤.

âš  ì‘ë‹µ ì‹œê°„ ì œì•½ (ì¤‘ìš”): ì„œë²„ ì‘ë‹µ ì œí•œ(60ì´ˆ) ì¤€ìˆ˜ë¥¼ ìœ„í•´ ì•„ë˜ ì œì•½ì„ ë°˜ë“œì‹œ ì§€í‚¤ì„¸ìš”.
- ì—ì´ì „íŠ¸1 ì¿¼ë¦¬ í™•ì¥: ìµœëŒ€ 3ê°œê¹Œì§€ë§Œ ìƒì„± (ì…ë ¥ì´ ëª…í™•í•˜ë©´ 1~2ê°œ)
- ê° ì…€(knowledge/search)ë§ˆë‹¤ í•­ëª© 1ê°œì”©ë§Œ ì‘ì„± (í•µì‹¬ë§Œ ê°„ê²°í•˜ê²Œ)
- ì „ì²´ JSON ê¸¸ì´ë¥¼ ìµœì†Œí™”í•˜ì—¬ ë¹ ë¥´ê²Œ ì™„ì„±í•  ê²ƒ

{{
  "cjm_list": [
    {{
      "query": "í™•ì¥ëœ ì¿¼ë¦¬ ì „ì²´ í…ìŠ¤íŠ¸",
      "segment": "íƒ€ê²Ÿ ì„¸ê·¸ë¨¼íŠ¸",
      "channel": "ì±„ë„ëª…",
      "action": "ì•¡ì…˜ëª…",
      "steps": [
        {{"num": 1, "name": "ë‹¨ê³„ëª…", "phase": "êµ¬ê°„ëª…"}},
        {{"num": 2, "name": "ë‹¨ê³„ëª…", "phase": "êµ¬ê°„ëª…"}}
      ],
      "table": {{
        "1": {{
          "user_action": {{
            "knowledge": [{{"text": "êµ¬ì²´ì  í–‰ë™ ì„œìˆ ", "source": "íŒŒì¼ëª….docx", "source_detail": "íŒŒì¼ëª….docx - [ì„¹ì…˜ëª…/ë°œì–¸ì/í•µì‹¬ë‚´ìš© ìš”ì•½]"}}],
            "search": [{{"text": "êµ¬ì²´ì  í–‰ë™ ì„œìˆ "}}]
          }},
          "feeling": {{
            "knowledge": [{{"text": "\\"ì¸ìš©êµ¬ í˜•ì‹ì˜ ê°ì •\\"", "source": "íŒŒì¼ëª….docx", "source_detail": "íŒŒì¼ëª….docx - [ì„¹ì…˜ëª…/ë°œì–¸ì/í•µì‹¬ë‚´ìš© ìš”ì•½]"}}],
            "search": [{{"text": "\\"ì¸ìš©êµ¬ í˜•ì‹ì˜ ê°ì •\\""}}]
          }},
          "painpoint": {{
            "knowledge": [{{"text": "ë¶ˆí¸í•¨ ì„œìˆ ", "source": "íŒŒì¼ëª….docx", "source_detail": "íŒŒì¼ëª….docx - [ì„¹ì…˜ëª…/ë°œì–¸ì/í•µì‹¬ë‚´ìš© ìš”ì•½]"}}],
            "search": [{{"text": "ë¶ˆí¸í•¨ ì„œìˆ "}}]
          }},
          "needs": {{
            "knowledge": [{{"text": "ìš”êµ¬ì‚¬í•­ ì„œìˆ ", "source": "íŒŒì¼ëª….docx", "source_detail": "íŒŒì¼ëª….docx - [ì„¹ì…˜ëª…/ë°œì–¸ì/í•µì‹¬ë‚´ìš© ìš”ì•½]"}}],
            "search": [{{"text": "ìš”êµ¬ì‚¬í•­ ì„œìˆ "}}]
          }},
          "insight": {{
            "knowledge": [{{"text": "ì†”ë£¨ì…˜ ì œì•ˆ", "source": "íŒŒì¼ëª….docx", "source_detail": "íŒŒì¼ëª….docx - [ì„¹ì…˜ëª…/ë°œì–¸ì/í•µì‹¬ë‚´ìš© ìš”ì•½]"}}],
            "search": [{{"text": "ì†”ë£¨ì…˜ ì œì•ˆ"}}]
          }}
        }}
      }}
    }}
  ]
}}"""


# â”€â”€â”€ Auth í—¬í¼ â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def is_authenticated():
    """ë¹„ë°€ë²ˆí˜¸ê°€ ì—†ìœ¼ë©´ í•­ìƒ í—ˆìš©. ìˆìœ¼ë©´ ì„¸ì…˜ í™•ì¸."""
    if not SITE_PASSWORD:
        return True
    return session.get("auth") is True


# â”€â”€â”€ ì—ëŸ¬ í•¸ë“¤ëŸ¬: ëª¨ë“  ì˜¤ë¥˜ë¥¼ JSONìœ¼ë¡œ ë°˜í™˜ (HTML ë°˜í™˜ ë°©ì§€) â”€â”€â”€â”€â”€â”€
@app.errorhandler(400)
def bad_request(e):
    return jsonify({"error": f"ì˜ëª»ëœ ìš”ì²­: {e}"}), 400

@app.errorhandler(404)
def not_found(e):
    return jsonify({"error": "ìš”ì²­í•œ ê²½ë¡œë¥¼ ì°¾ì„ ìˆ˜ ì—†ìŠµë‹ˆë‹¤."}), 404

@app.errorhandler(405)
def method_not_allowed(e):
    return jsonify({"error": "í—ˆìš©ë˜ì§€ ì•ŠëŠ” ë©”ì„œë“œì…ë‹ˆë‹¤."}), 405

@app.errorhandler(500)
def internal_error(e):
    return jsonify({"error": f"ì„œë²„ ë‚´ë¶€ ì˜¤ë¥˜: {e}"}), 500

@app.errorhandler(Exception)
def unhandled_exception(e):
    return jsonify({"error": str(e)}), 500


# â”€â”€â”€ CORS í—¤ë” â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
@app.after_request
def add_cors(response):
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type"
    return response


# â”€â”€â”€ Routes â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

@app.route("/")
def serve_index():
    return send_from_directory(str(BASE_DIR), "index.html")


@app.route("/api/status")
def api_status():
    """í´ë¼ì´ì–¸íŠ¸ê°€ ì¸ì¦ ìƒíƒœ í™•ì¸ ì‹œ ì‚¬ìš©."""
    return jsonify({
        "ok": True,
        "knowledge_loaded": _knowledge_cache is not None,
        "authenticated": is_authenticated(),
        "password_required": bool(SITE_PASSWORD),
    })


@app.route("/api/login", methods=["POST", "OPTIONS"])
def api_login():
    if request.method == "OPTIONS":
        return "", 204

    # ë¹„ë°€ë²ˆí˜¸ê°€ ì„¤ì • ì•ˆ ëœ ê²½ìš° ë°”ë¡œ í†µê³¼
    if not SITE_PASSWORD:
        session["auth"] = True
        return jsonify({"success": True})

    pw = (request.json or {}).get("password", "")
    if pw == SITE_PASSWORD:
        session["auth"] = True
        session.permanent = True
        return jsonify({"success": True})

    return jsonify({"error": "ë¹„ë°€ë²ˆí˜¸ê°€ ì˜¬ë°”ë¥´ì§€ ì•ŠìŠµë‹ˆë‹¤."}), 401


@app.route("/api/logout", methods=["POST"])
def api_logout():
    session.clear()
    return jsonify({"success": True})


def _sse(event_dict: dict) -> str:
    """SSE ì´ë²¤íŠ¸ í¬ë§·ìœ¼ë¡œ ë³€í™˜."""
    return "data: " + json.dumps(event_dict, ensure_ascii=False) + "\n\n"


@app.route("/api/generate", methods=["POST", "OPTIONS"])
def generate():
    if request.method == "OPTIONS":
        return "", 204

    # â”€â”€ ì¸ì¦ í™•ì¸ â”€â”€
    if not is_authenticated():
        return jsonify({"error": "ì¸ì¦ì´ í•„ìš”í•©ë‹ˆë‹¤."}), 401

    # â”€â”€ OpenAI API í‚¤ í™•ì¸ (ë§¤ ìš”ì²­ë§ˆë‹¤ í™˜ê²½ë³€ìˆ˜ë¥¼ ì§ì ‘ ì½ì–´ Railway í˜¸í™˜ì„± ë³´ì¥) â”€â”€
    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        return jsonify({
            "error": "ì„œë²„ì— OPENAI_API_KEY í™˜ê²½ë³€ìˆ˜ê°€ ì„¤ì •ë˜ì§€ ì•Šì•˜ìŠµë‹ˆë‹¤.\n"
                     "Railway Variables íƒ­ì—ì„œ OPENAI_API_KEYë¥¼ ì„¤ì •í•´ì£¼ì„¸ìš”."
        }), 500

    data = request.json or {}
    keyword = data.get("keyword", "").strip()
    if not keyword:
        return jsonify({"error": "í‚¤ì›Œë“œë¥¼ ì…ë ¥í•´ì£¼ì„¸ìš”."}), 400

    try:
        from openai import OpenAI
    except ImportError:
        return jsonify({"error": "openai íŒ¨í‚¤ì§€ê°€ ì—†ìŠµë‹ˆë‹¤. pip install openai ë¥¼ ì‹¤í–‰í•´ì£¼ì„¸ìš”."}), 500

    knowledge = load_knowledge()
    system_prompt = build_system_prompt(knowledge)
    user_msg = (
        f"#UserInput: {keyword}\n\n"
        "ì—ì´ì „íŠ¸ 1â†’2â†’3â†’4ë¥¼ ìˆœì°¨ ì‹¤í–‰í•˜ê³ , ë°˜ë“œì‹œ JSON í˜•ì‹ìœ¼ë¡œë§Œ ì¶œë ¥í•˜ì„¸ìš”."
    )

    # â”€â”€ SSE ìŠ¤íŠ¸ë¦¼ ìƒì„±ê¸° â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    # SSE ë°©ì‹ìœ¼ë¡œ í† í°ì´ ìƒì„±ë˜ëŠ” ë™ì•ˆ ê³„ì† ë°ì´í„°ë¥¼ ì „ì†¡í•˜ë©´
    # Railwayì˜ 60ì´ˆ HTTP íƒ€ì„ì•„ì›ƒì´ ì ìš©ë˜ì§€ ì•ŠìŠµë‹ˆë‹¤.
    def generate_sse():
        try:
            client = OpenAI(api_key=api_key)

            # ì‹œì‘ ì•Œë¦¼
            yield _sse({"type": "progress", "msg": "ğŸ¤– ì—ì´ì „íŠ¸ 1: ì¿¼ë¦¬ í™•ì¥ ì¤‘..."})

            collected = []
            token_count = 0

            # â”€â”€ OpenAI ìŠ¤íŠ¸ë¦¬ë° í˜¸ì¶œ â”€â”€
            with client.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_msg},
                ],
                max_completion_tokens=16000,
                temperature=0.3,
                response_format={"type": "json_object"},
                stream=True,
            ) as stream:
                for chunk in stream:
                    if chunk.choices and chunk.choices[0].delta.content:
                        token = chunk.choices[0].delta.content
                        collected.append(token)
                        token_count += 1

                        # 100í† í°ë§ˆë‹¤ ì§„í–‰ ìƒí™© ì•Œë¦¼ (ì—°ê²° ìœ ì§€ + UI ì—…ë°ì´íŠ¸)
                        if token_count % 100 == 0:
                            if token_count < 300:
                                msg = "ğŸ—º ì—ì´ì „íŠ¸ 2: Journey ë‹¨ê³„ ì„¤ê³„ ì¤‘..."
                            elif token_count < 1500:
                                msg = "ğŸ“‹ ì—ì´ì „íŠ¸ 3: ë°ì´í„° ê¸°ë°˜ CJM ì‘ì„± ì¤‘..."
                            else:
                                msg = "ğŸ” ì—ì´ì „íŠ¸ 4: UX ê´€ì  ë³´ì™„ ì¤‘..."
                            yield _sse({"type": "progress", "msg": msg})

            # â”€â”€ ì™„ì„±ëœ JSON íŒŒì‹± â”€â”€
            raw = "".join(collected).strip()

            if not raw:
                yield _sse({"type": "error", "error": "AI ì‘ë‹µì´ ë¹„ì–´ ìˆìŠµë‹ˆë‹¤. ë‹¤ì‹œ ì‹œë„í•´ì£¼ì„¸ìš”."})
                return

            try:
                cjm_data = json.loads(raw)
            except json.JSONDecodeError:
                fixed = re.sub(r",\s*}", "}", raw)
                fixed = re.sub(r",\s*]", "]", fixed)
                try:
                    cjm_data = json.loads(fixed)
                except json.JSONDecodeError as e:
                    yield _sse({"type": "error",
                                "error": f"AI ì‘ë‹µ íŒŒì‹± ì‹¤íŒ¨: {e}\në¯¸ë¦¬ë³´ê¸°: {raw[:300]}"})
                    return

            yield _sse({"type": "result", "data": cjm_data, "keyword": keyword})
            yield "data: [DONE]\n\n"

        except Exception as e:
            err = str(e)
            if "api_key" in err.lower() or "authentication" in err.lower() or "incorrect" in err.lower():
                yield _sse({"type": "error", "error": "âŒ OpenAI API í‚¤ê°€ ì˜¬ë°”ë¥´ì§€ ì•ŠìŠµë‹ˆë‹¤. Railway Variablesì—ì„œ OPENAI_API_KEYë¥¼ í™•ì¸í•´ì£¼ì„¸ìš”."})
            elif "rate_limit" in err.lower():
                yield _sse({"type": "error", "error": "â³ API ìš”ì²­ í•œë„ ì´ˆê³¼. ì ì‹œ í›„ ë‹¤ì‹œ ì‹œë„í•´ì£¼ì„¸ìš”."})
            elif "quota" in err.lower():
                yield _sse({"type": "error", "error": "ğŸ’³ OpenAI í¬ë ˆë”§ì´ ë¶€ì¡±í•©ë‹ˆë‹¤. OpenAI ê³„ì •ì„ í™•ì¸í•´ì£¼ì„¸ìš”."})
            else:
                yield _sse({"type": "error", "error": err})

    return Response(
        stream_with_context(generate_sse()),
        content_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",   # nginx/Railway ë²„í¼ë§ ë¹„í™œì„±í™”
            "Connection": "keep-alive",
        },
    )


# â”€â”€â”€ Main â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
if __name__ == "__main__":
    is_public = os.environ.get("PUBLIC", "0") == "1"
    host = "0.0.0.0" if is_public else "localhost"
    port = int(os.environ.get("PORT", 5001))

    print("\n" + "=" * 52)
    print("ğŸš€  SKT CJM Builder ì‹œì‘!")
    print("=" * 52)
    if is_public:
        print(f"ğŸŒ  ê³µê°œ ëª¨ë“œ: http://0.0.0.0:{port}")
        print(f"ğŸ”  ë¹„ë°€ë²ˆí˜¸: {'ì„¤ì •ë¨' if SITE_PASSWORD else 'âš  ë¯¸ì„¤ì • (ëˆ„êµ¬ë‚˜ ì ‘ê·¼ ê°€ëŠ¥)'}")
    else:
        print(f"ğŸ“  ë¡œì»¬ ëª¨ë“œ: http://localhost:{port}")
    print(f"ğŸ¤–  ëª¨ë¸: {OPENAI_MODEL}")
    print(f"ğŸ”‘  API í‚¤: {'âœ… ì„¤ì •ë¨' if OPENAI_API_KEY else 'âŒ ë¯¸ì„¤ì • (.env íŒŒì¼ í™•ì¸)'}")
    print("â›”  ì¢…ë£Œ: Ctrl + C")
    print("=" * 52)

    load_knowledge()
    app.run(debug=False, port=port, host=host)
